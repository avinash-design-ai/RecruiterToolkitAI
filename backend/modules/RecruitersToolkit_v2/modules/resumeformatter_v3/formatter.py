from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_COLOR_INDEX
FONT="Times New Roman"
SIZE=11
def fmt(p):
 p.alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
 pf=p.paragraph_format;pf.space_before=Pt(0);pf.space_after=Pt(0);pf.line_spacing=1
def runfmt(r,bold=False,italic=False,size=SIZE):
 r.font.name=FONT;r.font.size=Pt(size);r.bold=bold;r.italic=italic
def heading(doc,t):
 p=doc.add_paragraph();fmt(p);rr=p.add_run(t.upper());runfmt(rr,bold=True)
def placeholder(doc,t):
 p=doc.add_paragraph();fmt(p);rr=p.add_run(t);runfmt(rr);rr.font.highlight_color=WD_COLOR_INDEX.YELLOW
def bullets(doc,items):

 for i in items:
  p=doc.add_paragraph(style="List Bullet");fmt(p);rr=p.add_run(i);runfmt(rr)
def project_header(doc,prj):
 p=doc.add_paragraph();fmt(p)
 left=prj.client or "<< ADD CLIENT >>"
 if getattr(prj,"location",""): left+=", "+prj.location
 rr=p.add_run(left);runfmt(rr,bold=True)
 p.add_run("    ")
 rr=p.add_run(prj.role or "<< ADD ROLE >>");runfmt(rr,italic=True)
 rr=p.add_run(", "+(prj.duration or "<< ADD DURATION >>"));runfmt(rr,bold=True)
def generate_resume(resume,output_file):
 doc=Document()
 p=doc.add_paragraph();p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
 rr=p.add_run(resume.name or "<< ADD NAME >>");runfmt(rr,bold=True,size=12)
 p=doc.add_paragraph();p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
 contact=" | ".join([x for x in [getattr(resume,"phone",""),getattr(resume,"email",""),getattr(resume,"linkedin","")] if x]);rr=p.add_run(contact);runfmt(rr)
 heading(doc,"Professional Summary")
 bullets(doc,resume.summary) if getattr(resume,"summary",[]) else placeholder(doc,"<< ADD PROFESSIONAL SUMMARY >>")
 heading(doc,"Technical Skills")
 bullets(doc,resume.technical_skills) if getattr(resume,"technical_skills",[]) else placeholder(doc,"<< ADD TECHNICAL SKILLS >>")
 heading(doc,"Education")
 bullets(doc,resume.education) if getattr(resume,"education",[]) else placeholder(doc,"<< ADD EDUCATION >>")
 heading(doc,"Certifications")
 bullets(doc,resume.certifications) if getattr(resume,"certifications",[]) else placeholder(doc,"<< ADD CERTIFICATIONS >>")
 heading(doc,"Professional Experience")
 if getattr(resume,"projects",[]):
  for prj in resume.projects:
   project_header(doc,prj)
   if getattr(prj,"description",""):
    p=doc.add_paragraph();fmt(p);rr=p.add_run(prj.description);runfmt(rr)
   if getattr(prj,"responsibilities",[]): bullets(doc,prj.responsibilities)
   if getattr(prj,"environment",""):
    p=doc.add_paragraph();fmt(p);rr=p.add_run("Environment: ");runfmt(rr,bold=True);rr=p.add_run(prj.environment);runfmt(rr)
 else:
  placeholder(doc,"<< ADD PROFESSIONAL EXPERIENCE >>")
 if getattr(resume,"other_details",[]):
  heading(doc,"Other Details");bullets(doc,resume.other_details)
 doc.save(output_file)
