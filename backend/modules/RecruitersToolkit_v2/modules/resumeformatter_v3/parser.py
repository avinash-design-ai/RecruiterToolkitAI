from docx import Document
import pdfplumber


class ResumeReader:

    def read(self, file_path):

        if file_path.lower().endswith(".docx"):

            return self.read_docx(file_path)

        elif file_path.lower().endswith(".pdf"):

            return self.read_pdf(file_path)

        else:

            raise Exception("Unsupported file type.")

    # --------------------------------------------------
    # DOCX
    # --------------------------------------------------

    def read_docx(self, file_path):

        doc = Document(file_path)

        output = []

        # -----------------------------
        # Normal paragraphs
        # -----------------------------

        for para in doc.paragraphs:

            text = para.text.strip()

            if text:

                output.append(text)

        # -----------------------------
        # Tables
        # -----------------------------

        for table in doc.tables:

            for row in table.rows:

                cells = []

                for cell in row.cells:

                    value = cell.text.strip()

                    if value:

                        cells.append(value)

                if cells:

                    output.append(" | ".join(cells))

        return "\n".join(output)

    # --------------------------------------------------
    # PDF
    # --------------------------------------------------

    def read_pdf(self, file_path):

        output = []

        with pdfplumber.open(file_path) as pdf:

            for page in pdf.pages:

                text = page.extract_text()

                if text:

                    output.append(text)

        return "\n".join(output)
