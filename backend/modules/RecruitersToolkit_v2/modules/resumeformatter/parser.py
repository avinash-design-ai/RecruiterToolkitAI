from docx import Document
import fitz


def read_resume(file_path):
    """
    Reads DOCX or PDF resume and returns plain text.
    """

    if file_path.lower().endswith(".docx"):

        doc = Document(file_path)

        lines = []

        # Read paragraphs
        for para in doc.paragraphs:

            text = para.text.strip()

            if text:
                lines.append(text)

        # Read tables (many recruiter resumes use tables)
        for table in doc.tables:

            for row in table.rows:

                for cell in row.cells:

                    text = cell.text.strip()

                    if text:
                        lines.append(text)

        return "\n".join(lines)

    elif file_path.lower().endswith(".pdf"):

        pdf = fitz.open(file_path)

        text = ""

        for page in pdf:

            text += page.get_text()

        pdf.close()

        return text

    else:

        raise Exception(
            "Unsupported file format. Please upload DOCX or PDF."
        )
