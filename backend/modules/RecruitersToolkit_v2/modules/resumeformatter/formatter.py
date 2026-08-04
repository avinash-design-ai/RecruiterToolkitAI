from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX


# ---------------------------------------------
# HEADING
# ---------------------------------------------
def heading(doc, text):

    p = doc.add_paragraph()

    r = p.add_run(text)

    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)


# ---------------------------------------------
# BULLET
# ---------------------------------------------
def bullet(doc, text):

    p = doc.add_paragraph(style="List Bullet")

    r = p.add_run(text)

    r.font.name = "Times New Roman"
    r.font.size = Pt(11)


# ---------------------------------------------
# PLACEHOLDER
# ---------------------------------------------
def placeholder(doc, text):

    p = doc.add_paragraph()

    r = p.add_run(text)

    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW


# ---------------------------------------------
# CONTACT
# ---------------------------------------------
def add_contact(doc, resume):

    # Name

    p = doc.add_paragraph()

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    r = p.add_run(
        resume.name if resume.name else "PLEASE ADD NAME"
    )

    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    # Phone / Email

    p = doc.add_paragraph()

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    phone = resume.phone if resume.phone else "PLEASE ADD PHONE"

    email = resume.email if resume.email else "PLEASE ADD EMAIL"

    r = p.add_run(f"{phone} / {email}")

    r.font.name = "Times New Roman"
    r.font.size = Pt(11)


# ---------------------------------------------
# SUMMARY
# ---------------------------------------------
def add_summary(doc, resume):

    heading(doc, "PROFESSIONAL SUMMARY")

    if resume.summary:

        for item in resume.summary:

            bullet(doc, item)

    else:

        placeholder(doc, "PLEASE ADD SUMMARY")


# ---------------------------------------------
# EDUCATION
# ---------------------------------------------
def add_education(doc, resume):

    heading(doc, "EDUCATION")

    if resume.education:

        for item in resume.education:

            bullet(doc, item)

    else:

        placeholder(doc, "PLEASE ADD EDUCATION")


# ---------------------------------------------
# TECHNICAL SKILLS
# ---------------------------------------------
def add_skills(doc, resume):

    heading(doc, "TECHNICAL SKILLS")

    if resume.technical_skills:

        for item in resume.technical_skills:

            bullet(doc, item)

    else:

        placeholder(doc, "PLEASE ADD TECHNICAL SKILLS")


# ---------------------------------------------
# CERTIFICATIONS
# ---------------------------------------------
def add_certifications(doc, resume):

    heading(doc, "CERTIFICATIONS")

    if resume.certifications:

        for item in resume.certifications:

            bullet(doc, item)

    else:

        placeholder(doc, "PLEASE ADD CERTIFICATIONS")


# ---------------------------------------------
# EXPERIENCE
# ---------------------------------------------
def add_experience(doc):

    heading(doc, "PROFESSIONAL EXPERIENCE")

    placeholder(doc, "PLEASE ADD EXPERIENCE")


# ---------------------------------------------
# MAIN
# ---------------------------------------------
def generate_resume(resume, output_file):

    doc = Document()

    add_contact(doc, resume)

    add_summary(doc, resume)

    add_education(doc, resume)

    add_skills(doc, resume)

    add_certifications(doc, resume)

    add_experience(doc)

    doc.save(output_file)
