from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn


FONT = "Times New Roman"
SIZE = 11


# ---------------------------------------------------------
# Font
# ---------------------------------------------------------

def style_run(
    run,
    bold=False,
    italic=False,
    highlight=False
):

    run.font.name = FONT
    run.font.size = Pt(SIZE)
    run.bold = bold
    run.italic = italic

    if run._element.rPr is not None:

        run._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            FONT
        )

    if highlight:

        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


# ---------------------------------------------------------
# Clear Paragraph
# ---------------------------------------------------------

def clear_paragraph(paragraph):

    element = paragraph._element

    while len(element):

        element.remove(element[0])


# ---------------------------------------------------------
# Find Placeholder
# ---------------------------------------------------------

def find_placeholder(
    doc,
    placeholder
):

    for paragraph in doc.paragraphs:

        if placeholder in paragraph.text:

            return paragraph

    return None


# ---------------------------------------------------------
# Replace Single Text
# ---------------------------------------------------------

def replace_text(
    doc,
    placeholder,
    value,
    highlight=False,
    bold=False,
    size=SIZE
):
    for p in doc.paragraphs:

        if placeholder not in p.text:

            continue

        # Preserve the paragraph alignment from the template
        alignment = p.alignment

        full_text = p.text.replace(

            placeholder,

            value

        )

        clear_paragraph(p)

        p.alignment = alignment

        run = p.add_run(full_text)

        style_run(
            run,
            bold=bold,
            highlight=highlight
        )

        run.font.size = Pt(size)

        return


# ---------------------------------------------------------
# Replace Bullet List
# ---------------------------------------------------------

def replace_bullets(
    doc,
    placeholder,
    values,
    empty_text
):

    p = find_placeholder(
        doc,
        placeholder
    )

    if p is None:

        return

    clear_paragraph(p)

    if not values:

        run = p.add_run(
            empty_text
        )

        style_run(
            run,
            highlight=True
        )

        return

    for item in values:

        item = " ".join(
            str(item).split()
        )

        if not item:

            continue

        run = p.add_run(
            "• " + item
        )

        style_run(run)

        p.add_run("\n")


# ---------------------------------------------------------
# Contact
# ---------------------------------------------------------

def write_contact(
    doc,
    resume
):

    replace_text(

        doc,

        "{{NAME}}",

        resume.name
        if resume.name
        else "ADD NAME",

        highlight=not bool(resume.name),

        bold=True,

        size=12

    )

    replace_text(

        doc,

        "{{PHONE}}",

        resume.phone
        if resume.phone
        else "ADD PHONE",

        not bool(resume.phone)

    )

    replace_text(

        doc,

        "{{EMAIL}}",

        resume.email
        if resume.email
        else "ADD EMAIL",

        not bool(resume.email)

    )

    replace_text(

        doc,

        "{{LINKEDIN}}",

        resume.linkedin

    )


# ---------------------------------------------------------
# Summary
# ---------------------------------------------------------

def write_summary(
    doc,
    resume
):

    replace_bullets(

        doc,

        "{{SUMMARY}}",

        resume.summary,

        "ADD PROFESSIONAL SUMMARY"

    )


# ---------------------------------------------------------
# Education
# ---------------------------------------------------------

def write_education(
    doc,
    resume
):

    replace_bullets(

        doc,

        "{{EDUCATION}}",

        resume.education,

        "ADD EDUCATION"

    )


# ---------------------------------------------------------
# Certifications
# ---------------------------------------------------------

def write_certifications(
    doc,
    resume
):

    replace_bullets(

        doc,

        "{{CERTIFICATIONS}}",

        resume.certifications,

        ""

    )


# ---------------------------------------------------------
# Skills
# ---------------------------------------------------------

def write_skills(
    doc,
    resume
):

    p = find_placeholder(
        doc,
        "{{SKILLS}}"
    )

    if p is None:

        return

    clear_paragraph(p)

    if not resume.technical_skills:

        run = p.add_run(
            "ADD TECHNICAL SKILLS"
        )

        style_run(
            run,
            highlight=True
        )

        return

    for category, values in resume.technical_skills.items():

        run = p.add_run(
            category + " : "
        )

        style_run(
            run,
            bold=True
        )

        run = p.add_run(
            ", ".join(values)
        )

        style_run(run)

        p.add_run("\n")

# ---------------------------------------------------------
# Experience
# ---------------------------------------------------------

def write_experience(
    doc,
    resume
):

    p = find_placeholder(
        doc,
        "{{EXPERIENCE}}"
    )

    if p is None:
        return

    clear_paragraph(p)

    if not resume.experience:

        run = p.add_run(
            "ADD EXPERIENCE"
        )

        style_run(
            run,
            highlight=True
        )

        return

    for job in resume.experience:

        # -----------------------------------
        # Header
        # -----------------------------------

        first = True

        def separator():
            nonlocal first

            if not first:

                run = p.add_run(", ")

                style_run(run)

            first = False

        # Client

                # Client / Employer

        separator()

        if job.employer and job.client:

            display_client = (
                f"{job.employer}/{job.client}"
            )

        elif job.employer:

            display_client = job.employer

        elif job.client:

            display_client = job.client

        else:

            display_client = "ADD CLIENT"

        run = p.add_run(
            display_client
        )

        if display_client == "ADD CLIENT":

            style_run(
                run,
                bold=True,
                highlight=True
            )

        else:

            style_run(
                run,
                bold=True
            )

        # Location

        if job.location:

            separator()

            run = p.add_run(
                job.location
            )

            style_run(
                run,
                bold=True
            )

        # Role

        separator()

        if job.role:

            run = p.add_run(
                job.role
            )

            style_run(
                run,
                italic=True
            )

        else:

            run = p.add_run(
                "ADD ROLE"
            )

            style_run(
                run,
                italic=True,
                highlight=True
            )

        # Duration

        separator()

        if job.duration:

            run = p.add_run(
                job.duration
            )

            style_run(
                run,
                bold=True
            )

        else:

            run = p.add_run(
                "ADD DURATION"
            )

            style_run(
                run,
                bold=True,
                highlight=True
            )

        p.add_run("\n")

        # -----------------------------------
        # Responsibilities
        # -----------------------------------

        for project in job.projects:

            for bullet in project.responsibilities:

                bullet = " ".join(
                    str(bullet).split()
                )

                if not bullet:

                    continue

                run = p.add_run(
                    "• " + bullet
                )

                style_run(run)

                p.add_run("\n")

        # -----------------------------------
        # Environment
        # -----------------------------------

        run = p.add_run(
            "Environment : "
        )

        style_run(
            run,
            bold=True
        )

        if job.environment:

            env = []

            for item in job.environment:

                item = " ".join(
                    str(item).split()
                )

                if item:

                    env.append(item)

            run = p.add_run(
                ", ".join(env)
            )

            style_run(run)

        else:

            run = p.add_run(
                "ADD ENVIRONMENT"
            )

            style_run(
                run,
                highlight=True
            )

        p.add_run("\n\n")


# ---------------------------------------------------------
# Generate Resume
# ---------------------------------------------------------

def generate_resume(
    resume,
    template_file,
    output_file
):

    doc = Document(
        template_file
    )

    write_contact(
        doc,
        resume
    )

    write_summary(
        doc,
        resume
    )

    write_education(
        doc,
        resume
    )

    write_certifications(
        doc,
        resume
    )

    write_skills(
        doc,
        resume
    )

    write_experience(
        doc,
        resume
    )

    doc.save(
        output_file
    )

    return output_file
